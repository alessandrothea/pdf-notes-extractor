#!/usr/bin/env python3
"""
PDF Annotation Extractor
Extracts annotations (highlights, comments, notes, etc.) from PDF files
and exports them to Markdown, Plain Text, Microsoft Word, or PDF format.

Usage:
    python extract_pdf_annotations.py input.pdf [options]

Requirements:
    pip install pypdf python-docx fpdf2 click rich
"""

import html
import click
from pathlib import Path
from datetime import datetime

from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.text import Text
from rich.theme import Theme
from rich.rule import Rule
from rich import box

# ── Rich console setup ─────────────────────────────────────────────────────────

THEME = Theme({
    "info":    "bold cyan",
    "success": "bold green",
    "warning": "bold yellow",
    "error":   "bold red",
    "page":    "bold magenta",
    "label":   "dim cyan",
    "path":    "underline bright_blue",
    "count":   "bold yellow",
    "author":  "italic bright_cyan",
    "preview": "dim white",
    "type":    "bold white",
})

console = Console(theme=THEME)


# ── Annotation type map ────────────────────────────────────────────────────────

ANNOTATION_TYPES = {
    "/Text":           "Note",
    "/FreeText":       "Free Text",
    "/Line":           "Line",
    "/Square":         "Rectangle",
    "/Circle":         "Ellipse",
    "/Polygon":        "Polygon",
    "/PolyLine":       "Polyline",
    "/Highlight":      "Highlight",
    "/Underline":      "Underline",
    "/Squiggly":       "Squiggly Underline",
    "/StrikeOut":      "Strikethrough",
    "/Stamp":          "Stamp",
    "/Caret":          "Caret",
    "/Ink":            "Ink",
    "/Popup":          "Popup",
    "/FileAttachment": "File Attachment",
    "/Sound":          "Sound",
    "/Movie":          "Movie",
    "/Widget":         "Form Field",
    "/Screen":         "Screen",
    "/PrinterMark":    "Printer Mark",
    "/TrapNet":        "Trap Network",
    "/Watermark":      "Watermark",
    "/3D":             "3D",
    "/Redact":         "Redaction",
    "/Link":           "Link",
}


# ── Core extractor ─────────────────────────────────────────────────────────────

def _extract_quad_text(page, quad_points) -> str:
    """Extract the text covered by QuadPoints on a page."""
    try:
        quads = [quad_points[i:i+8] for i in range(0, len(quad_points), 8)]
        regions = []
        for quad in quads:
            xs = [float(quad[i]) for i in range(0, 8, 2)]
            ys = [float(quad[i]) for i in range(1, 8, 2)]
            # small tolerance for floating-point mismatches
            regions.append((min(xs) - 1, min(ys) - 1, max(xs) + 1, max(ys) + 1))

        parts: list[str] = []

        def visitor(*args):
            text, cm, tm = args[0], args[1], args[2]
            if not text or not text.strip():
                return
            # Combine graphics-state matrix (cm) with text matrix (tm)
            # to get actual page coordinates.
            x = cm[0] * tm[4] + cm[2] * tm[5] + cm[4]
            y = cm[1] * tm[4] + cm[3] * tm[5] + cm[5]
            for x0, y0, x1, y1 in regions:
                if x0 <= x <= x1 and y0 <= y <= y1:
                    parts.append(text)
                    break

        page.extract_text(visitor_text=visitor)
        return "".join(parts).strip()
    except Exception:
        return ""


def _build_section_map(reader) -> list[tuple[int, str]]:
    """Return a sorted list of (1-based page number, section title) from the PDF outline."""
    sections: list[tuple[int, str]] = []

    def traverse(outline):
        for item in outline:
            if isinstance(item, list):
                traverse(item)
            else:
                try:
                    page_num = reader.get_destination_page_number(item) + 1
                    sections.append((page_num, item.title))
                except Exception:
                    pass

    try:
        traverse(reader.outline)
    except Exception:
        pass

    return sorted(sections)


def _section_for_page(sections: list[tuple[int, str]], page: int) -> str:
    """Return the title of the last section that starts on or before the given page."""
    result = ""
    for page_num, title in sections:
        if page_num <= page:
            result = title
        else:
            break
    return result


def extract_annotations(pdf_path: Path) -> list[dict]:
    """Extract all annotations from a PDF, page by page."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise click.ClickException("pypdf is required.  Run:  pip install pypdf")

    reader = PdfReader(str(pdf_path))
    sections = _build_section_map(reader)
    annotations = []

    for page_num, page in enumerate(reader.pages, start=1):
        if "/Annots" not in page:
            continue

        for annot_ref in page["/Annots"]:
            try:
                annot = annot_ref.get_object()
            except Exception:
                continue

            subtype = annot.get("/Subtype", "")
            # Skip invisible / structural annotations
            if subtype in ("/Link", "/Widget", "/Popup", "/PrinterMark", "/TrapNet"):
                continue

            content = annot.get("/Contents", "").strip() if annot.get("/Contents") else ""
            author  = annot.get("/T", "").strip()        if annot.get("/T")         else ""
            subject = annot.get("/Subj", "").strip()     if annot.get("/Subj")      else ""

            marked_text = ""
            if "/QuadPoints" in annot:
                marked_text = _extract_quad_text(page, annot["/QuadPoints"])

            raw_date = annot.get("/M", "")
            date_str = _parse_pdf_date(str(raw_date)) if raw_date else ""

            color_str = ""
            if "/C" in annot:
                try:
                    c = annot["/C"]
                    if len(c) == 3:
                        color_str = "#{:02X}{:02X}{:02X}".format(
                            int(c[0] * 255), int(c[1] * 255), int(c[2] * 255)
                        )
                except Exception:
                    pass

            annotations.append({
                "page":        page_num,
                "section":     _section_for_page(sections, page_num),
                "type":        ANNOTATION_TYPES.get(str(subtype), str(subtype).lstrip("/")),
                "author":      author,
                "date":        date_str,
                "subject":     subject,
                "content":     content,
                "marked_text": marked_text,
                "color":       color_str,
            })

    return annotations


def _parse_pdf_date(raw: str) -> str:
    """Convert a PDF date string (D:YYYYMMDDHHmmSS...) to a readable format."""
    raw = raw.strip().lstrip("D:").replace("'", "")
    for fmt in ("%Y%m%d%H%M%S%z", "%Y%m%d%H%M%S", "%Y%m%d%H%M", "%Y%m%d"):
        try:
            return datetime.strptime(
                raw[: len(fmt.replace("%z", ""))], fmt.replace("%z", "")
            ).strftime("%Y-%m-%d %H:%M")
        except ValueError:
            continue
    return raw


# ── File exporters ─────────────────────────────────────────────────────────────

def _group_by_page(annotations: list[dict]) -> dict[int, list[dict]]:
    groups: dict[int, list[dict]] = {}
    for a in annotations:
        groups.setdefault(a["page"], []).append(a)
    return groups


def export_markdown(annotations: list[dict], pdf_name: str, out_path: Path) -> None:
    lines = [
        f"# PDF Annotations — {pdf_name}",
        "",
        f"*Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}*  ",
        f"*Total annotations: {len(annotations)}*",
        "",
    ]

    for page, items in _group_by_page(annotations).items():
        lines.append(f"## Page {page}")
        lines.append("")
        for a in items:
            lines.append(f"### {a['type']}")
            meta = []
            if a["author"]: meta.append(f"**Author:** {a['author']}")
            if a["date"]:   meta.append(f"**Date:** {a['date']}")
            if meta: lines.append("- " + " · ".join(meta))
            if a["subject"]: lines.append(f"- **Subject:** {a['subject']}")
            if a["marked_text"]:
                if a["color"]:
                    lines.append(f'<mark style="background-color: {a["color"]}">{a["marked_text"]}</mark>')
                else:
                    lines.append(f"> {a['marked_text']}")
            if a["content"]:
                lines.append("- **Comment:**")
                lines.append(f"  > {a['content'].replace(chr(10), chr(10) + '  > ')}")
            lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    console.print(f"  [success]✓[/success]  Markdown   [dim]→[/dim] [path]{out_path}[/path]")


def export_plain_text(annotations: list[dict], pdf_name: str, out_path: Path) -> None:
    sep = "-" * 60
    lines = [
        f"PDF ANNOTATIONS — {pdf_name}",
        f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Total annotations: {len(annotations)}",
        sep, "",
    ]

    for page, items in _group_by_page(annotations).items():
        lines.append(f"PAGE {page}")
        lines.append(sep)
        for a in items:
            lines.append(f"Type    : {a['type']}")
            meta = "  ·  ".join(filter(None, [
                f"Author: {a['author']}" if a["author"] else "",
                f"Date: {a['date']}"     if a["date"]   else "",
            ]))
            if meta: lines.append(meta)
            if a["subject"]:     lines.append(f"Subject : {a['subject']}")
            if a["marked_text"]: lines.append(f"Marked  : {a['marked_text']}")
            if a["content"]:     lines.append(f"Comment : {a['content']}")
            lines.append("")
        lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    console.print(f"  [success]✓[/success]  Plain text [dim]→[/dim] [path]{out_path}[/path]")


def _word_highlight_run(run, hex_color: str) -> None:
    """Apply a custom background shading color to a Word run via XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    rPr.append(shd)


_FONT_TITLE = "Karla"
_FONT_BODY  = "Lato"


def _word_font(paragraph, name: str) -> None:
    """Set the font name on every run in a paragraph."""
    for run in paragraph.runs:
        run.font.name = name


def export_word(annotations: list[dict], pdf_name: str, out_path: Path) -> None:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise click.ClickException("python-docx is required.  Run:  pip install python-docx")

    doc = DocxDocument()

    # Default body font
    doc.styles["Normal"].font.name = _FONT_BODY

    title = doc.add_heading(f"PDF Annotations — {pdf_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _word_font(title, _FONT_TITLE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
        f"Total annotations: {len(annotations)}"
    )
    run.font.name = _FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    for page, items in _group_by_page(annotations).items():
        h1 = doc.add_heading(f"Page {page}", level=1)
        h1.paragraph_format.space_before = Pt(10)
        h1.paragraph_format.space_after  = Pt(2)
        _word_font(h1, _FONT_TITLE)

        for a in items:
            h2 = doc.add_heading(a["type"], level=2)
            h2.paragraph_format.space_before = Pt(6)
            h2.paragraph_format.space_after  = Pt(2)
            _word_font(h2, _FONT_TITLE)
            for run in h2.runs:
                run.font.size = Pt(9)

            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            def _compact(row):
                for cell in row.cells:
                    fmt = cell.paragraphs[0].paragraph_format
                    fmt.space_before = Pt(1)
                    fmt.space_after  = Pt(1)

            def _add_row(label, value):
                if not value:
                    return
                row = table.add_row()
                _compact(row)
                lbl = row.cells[0].paragraphs[0].add_run(label)
                lbl.bold = True
                lbl.font.name = _FONT_BODY
                lbl.font.size = Pt(9)
                val = row.cells[1].paragraphs[0].add_run(value)
                val.font.name = _FONT_BODY
                val.font.size = Pt(9)
                row.cells[0].width = 1_200_000
                row.cells[1].width = 6_500_000

            def _add_full_row(value, bold=False, highlight=None):
                if not value:
                    return
                row = table.add_row()
                merged = row.cells[0].merge(row.cells[1])
                merged.paragraphs[0].paragraph_format.space_before = Pt(1)
                merged.paragraphs[0].paragraph_format.space_after  = Pt(1)
                run = merged.paragraphs[0].add_run(value)
                run.bold = bold
                run.font.name = _FONT_BODY
                run.font.size = Pt(9)
                if highlight:
                    _word_highlight_run(run, highlight)

            # Author · Date: full-width, bold
            author_date = "  ·  ".join(filter(None, [a["author"], a["date"]]))
            _add_full_row(author_date, bold=True)
            _add_row("Subject", a["subject"])

            # Marked text: full-width, with highlight color
            _add_full_row(
                a["marked_text"],
                highlight=a["color"] if a["color"] else None,
            )

            _add_row("Comment", a["content"])

    doc.save(str(out_path))
    console.print(f"  [success]✓[/success]  Word doc   [dim]→[/dim] [path]{out_path}[/path]")


# ── Rich terminal markdown view ───────────────────────────────────────────────


def print_markdown_rich(annotations: list[dict], pdf_name: str) -> None:
    """Render annotations as rich output in the terminal."""
    console.print()
    console.print(f"[bold]{pdf_name}[/bold]")
    console.print(f"[dim]{datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  {len(annotations)} annotation(s)[/dim]")

    for page, items in _group_by_page(annotations).items():
        console.print()
        console.print(Rule(f"[bold]Page {page}[/bold]", style="bright_black"))

        for a in items:
            outer_rows = []

            if a["section"]:
                outer_rows.append(Text.from_markup(
                    f"[dim]§[/dim] [italic]{a['section']}[/italic]"
                ))

            if a["marked_text"]:
                if a["color"]:
                    outer_rows.append(Text.from_markup(
                        f"[on {a['color']}]{a['marked_text']}[/on {a['color']}]"
                    ))
                else:
                    outer_rows.append(Text.from_markup(f"[italic]{a['marked_text']}[/italic]"))

            has_inner = a["subject"] or a["content"]
            if has_inner:
                inner = Table(show_header=False, padding=(0, 0), expand=True)
                inner.add_column("label", no_wrap=True, max_width=20)
                inner.add_column("value", overflow="fold")
                inner_rows = []
                if a["subject"]:
                    inner_rows.append(("[dim]Subject[/dim]", a["subject"]))
                if a["content"]:
                    lbl_parts = []
                    if a["author"]: lbl_parts.append(f"[bold bright_cyan]{a['author']}[/bold bright_cyan]")
                    if a["date"]:   lbl_parts.append(f"[dim]{a['date']}[/dim]")
                    inner_rows.append(("[dim]Comment[/dim]", a["content"]))
                for i, (lbl, val) in enumerate(inner_rows):
                    inner.add_row(lbl, val, end_section=(i < len(inner_rows) - 1))
                outer_rows.append(inner)

            if not outer_rows:
                continue

            outer = Table(box=box.MINIMAL, show_header=False, padding=(0, 0), expand=True)
            outer.add_column("content", overflow="fold")
            for i, r in enumerate(outer_rows):
                # outer.add_row(r, end_section=(i < len(outer_rows) - 1))
                outer.add_row(r)

            title_parts = [f"[bold]{a['type']}[/bold]"]
            if a["author"]: title_parts.append(f"[bright_cyan]{a['author']}[/bright_cyan]")
            if a["date"]:   title_parts.append(f"[dim]{a['date']}[/dim]")

            console.print(Panel(
                outer,
                title="  [dim]·[/dim]  ".join(title_parts),
                title_align="left",
                border_style="bright_black",
                expand=False,
                padding=(0, 0),
            ))
            console.print()


# ── Rich terminal list view ────────────────────────────────────────────────────

def print_summary_table(annotations: list[dict]) -> None:
    """Render all annotations as a Rich table in the terminal."""
    table = Table(
        title="[bold]PDF Annotations[/bold]",
        box=box.ROUNDED,
        show_header=True,
        header_style="bold cyan",
        border_style="bright_black",
        expand=False,
        pad_edge=True,
        show_lines=True,
    )

    table.add_column("Page",    style="magenta",     justify="center", no_wrap=True)
    table.add_column("Type",    style="bold white",  no_wrap=True)
    table.add_column("Author",  style="bright_cyan", no_wrap=True)
    table.add_column("Date",    style="dim white",   no_wrap=True)
    table.add_column("Comment / Content", style="white", overflow="fold", min_width=32)

    for a in annotations:
        if a["content"]:
            preview = a["content"][:80] + ("…" if len(a["content"]) > 80 else "")
        elif a["marked_text"]:
            preview = f"[dim italic]{a['marked_text']}[/dim italic]"
        else:
            preview = "[dim]—[/dim]"

        table.add_row(
            str(a["page"]),
            a["type"],
            a["author"] or "[dim]—[/dim]",
            a["date"]   or "[dim]—[/dim]",
            preview,
        )

    console.print()
    console.print(table)
    console.print()


# ── PDF exporter ───────────────────────────────────────────────────────────────

def _find_font_file(*name_parts: str) -> str | None:
    """Search common system font directories for a .ttf whose stem contains all name_parts."""
    dirs = [
        Path.home() / "Library/Fonts",
        Path("/Library/Fonts"),
        Path.home() / ".fonts",
        Path("/usr/share/fonts"),
        Path("C:/Windows/Fonts"),
    ]
    for d in dirs:
        if not d.exists():
            continue
        for f in d.rglob("*.ttf"):
            stem = f.stem.lower()
            if all(p.lower() in stem for p in name_parts):
                return str(f)
    return None


def export_pdf(annotations: list[dict], pdf_name: str, out_path: Path) -> None:
    try:
        from fpdf import FPDF
        from fpdf.fonts import FontFace
    except ImportError:
        raise click.ClickException("fpdf2 is required.  Run:  pip install fpdf2")

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(True, margin=20)

    # Register fonts; fall back to a system Unicode font, then helvetica
    def _register(family: str, reg_parts: list[str], bld_parts: list[str]) -> str:
        reg = _find_font_file(*reg_parts)
        if not reg:
            return ""
        bld = _find_font_file(*bld_parts)
        try:
            pdf.add_font(family, style="",  fname=reg)
            pdf.add_font(family, style="B", fname=bld or reg)
            return family
        except Exception:
            return ""

    def _unicode_fallback() -> str:
        """Find and register a Unicode-capable system font to use when custom fonts are absent."""
        for name, parts in [
            ("Arial",          ["Arial"]),
            ("NotoSans",       ["NotoSans", "Regular"]),
            ("LiberationSans", ["Liberation", "Sans", "Regular"]),
            ("DejaVuSans",     ["DejaVuSans"]),
        ]:
            reg = _find_font_file(*parts)
            if reg:
                try:
                    pdf.add_font(name, style="",  fname=reg)
                    pdf.add_font(name, style="B", fname=reg)
                    return name
                except Exception:
                    continue
        return "helvetica"

    title_font = _register("Karla", ["Karla", "Regular"], ["Karla", "Bold"])
    body_font  = _register("Lato",  ["Lato",  "Regular"], ["Lato",  "Bold"])
    _fallback  = None  # resolved lazily, at most once

    def fallback() -> str:
        nonlocal _fallback
        if _fallback is None:
            _fallback = _unicode_fallback()
        return _fallback

    if not title_font:
        title_font = fallback()
    if not body_font:
        body_font = fallback()

    def hex_rgb(hex_color: str) -> tuple[int, int, int]:
        h = hex_color.lstrip("#")
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    LABEL_W = 33
    VALUE_W = 137
    GREY    = (245, 245, 245)

    pdf.add_page()

    # Document title
    pdf.set_font(title_font, "B", 18)
    pdf.cell(0, 10, f"PDF Annotations - {pdf_name}", align="C",
             new_x="LMARGIN", new_y="NEXT")

    # Meta line
    pdf.set_font(body_font, "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6,
             f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
             f"Total annotations: {len(annotations)}",
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    for page, items in _group_by_page(annotations).items():
        pdf.ln(8)
        pdf.set_font(title_font, "B", 13)
        pdf.cell(0, 7, f"Page {page}", new_x="LMARGIN", new_y="NEXT")

        for a in items:
            pdf.ln(2)
            pdf.set_font(title_font, "B", 9)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 5, a["type"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

            author_date = "  \u00b7  ".join(filter(None, [a["author"], a["date"]]))
            has_rows = any([author_date, a["subject"], a["marked_text"], a["content"]])
            if not has_rows:
                continue

            with pdf.table(
                col_widths=(LABEL_W, VALUE_W),
                borders_layout="ALL",
                line_height=5,
                text_align="LEFT",
            ) as table:
                if author_date:
                    row = table.row()
                    row.cell(author_date, colspan=2,
                             style=FontFace(font_family=body_font, font_style="B", font_size_pt=9))

                if a["subject"]:
                    row = table.row()
                    row.cell("Subject",   style=FontFace(font_family=body_font, font_style="B",
                                                         font_size_pt=9, fill_color=GREY))
                    row.cell(a["subject"], style=FontFace(font_family=body_font, font_size_pt=9))

                if a["marked_text"]:
                    fill = hex_rgb(a["color"]) if a["color"] else None
                    row = table.row()
                    style = FontFace(font_family=body_font, font_size_pt=9,
                                     fill_color=fill if fill else (255, 255, 255))
                    row.cell(a["marked_text"], colspan=2, style=style)

                if a["content"]:
                    row = table.row()
                    row.cell("Comment",    style=FontFace(font_family=body_font, font_style="B",
                                                          font_size_pt=9, fill_color=GREY))
                    row.cell(a["content"], style=FontFace(font_family=body_font, font_size_pt=9))

    pdf.output(str(out_path))
    console.print(f"  [success]✓[/success]  PDF        [dim]→[/dim] [path]{out_path}[/path]")


# ── CLI ────────────────────────────────────────────────────────────────────────

FORMAT_CHOICES = click.Choice(
    ["markdown", "md", "txt", "text", "docx", "word", "pdf", "all"],
    case_sensitive=False,
)


@click.command(context_settings={"help_option_names": ["-h", "--help"]})
@click.argument("pdf", type=click.Path(exists=True, dir_okay=False, path_type=Path))
@click.option(
    "--format", "-f", "fmt",
    type=FORMAT_CHOICES,
    default="markdown",
    show_default=True,
    help="Output format when saving to file (requires --output or --save).",
)
@click.option(
    "--output", "-o",
    type=click.Path(dir_okay=False, path_type=Path),
    default=None,
    help="Save to this file path instead of printing to screen.",
)
@click.option(
    "--save", "-s",
    is_flag=True,
    default=False,
    help="Save to a file next to the PDF (auto-named) instead of printing to screen.",
)
@click.option(
    "--list-only", "-l",
    is_flag=True,
    default=False,
    help="Print a rich table instead of rendered markdown.",
)
def main(pdf: Path, fmt: str, output: Path | None, save: bool, list_only: bool) -> None:
    """Extract annotations from a PDF and print them or save to a file.

    \b
    By default, annotations are rendered as markdown in the terminal.
    Use --output or --save to write to a file instead.

    \b
    Examples:
      extract_pdf_annotations.py report.pdf
      extract_pdf_annotations.py report.pdf --list-only
      extract_pdf_annotations.py report.pdf --save
      extract_pdf_annotations.py report.pdf --save -f txt
      extract_pdf_annotations.py report.pdf -f docx -o my_notes.docx
      extract_pdf_annotations.py report.pdf --save -f all
    """
    if pdf.suffix.lower() != ".pdf":
        console.print(
            f"[warning]⚠  '{pdf.name}' does not have a .pdf extension – proceeding anyway.[/warning]"
        )

    console.print(Panel(
        f"[info]Reading:[/info] [path]{pdf}[/path]",
        box=box.ROUNDED,
        border_style="bright_black",
        expand=False,
    ))

    annotations = extract_annotations(pdf)

    if not annotations:
        console.print("[warning]⚠  No annotations found in this PDF.[/warning]")
        return

    pages = len({a["page"] for a in annotations})
    console.print(
        f"[success]✓[/success]  Found [count]{len(annotations)}[/count] annotation(s) "
        f"across [count]{pages}[/count] page(s).\n"
    )

    if list_only:
        print_summary_table(annotations)
        return

    # Default: render to terminal as rich markdown
    if not output and not save:
        print_markdown_rich(annotations, pdf.name)
        return

    fmt = fmt.lower()

    def _resolve(suffix: str) -> Path:
        if output and fmt != "all":
            return output.with_suffix(suffix) if output.suffix.lower() != suffix else output
        return pdf.parent / f"{pdf.stem}_annotations{suffix}"

    console.print("[info]Exporting…[/info]")

    if fmt in ("markdown", "md", "all"):
        export_markdown(annotations, pdf.name, _resolve(".md"))

    if fmt in ("txt", "text", "all"):
        export_plain_text(annotations, pdf.name, _resolve(".txt"))

    if fmt in ("docx", "word", "all"):
        export_word(annotations, pdf.name, _resolve(".docx"))

    if fmt in ("pdf", "all"):
        export_pdf(annotations, pdf.name, _resolve(".pdf"))

    console.print("\n[success]Done![/success]")


if __name__ == "__main__":
    main()
