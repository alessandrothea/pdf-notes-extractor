from extract_pdf_annotations import (
    export_markdown,
    export_plain_text,
    export_word,
    export_pdf,
)


# ── export_markdown ────────────────────────────────────────────────────────────

def test_markdown_creates_file(sample_annotations, tmp_path):
    out = tmp_path / "out.md"
    export_markdown(sample_annotations, "test.pdf", out)
    assert out.exists() and out.stat().st_size > 0


def test_markdown_contains_type_heading(sample_annotations, tmp_path):
    out = tmp_path / "out.md"
    export_markdown(sample_annotations, "test.pdf", out)
    assert "### Highlight" in out.read_text()


def test_markdown_contains_marked_text(sample_annotations, tmp_path):
    out = tmp_path / "out.md"
    export_markdown(sample_annotations, "test.pdf", out)
    assert "brown fox" in out.read_text()


def test_markdown_contains_comment(sample_annotations, tmp_path):
    out = tmp_path / "out.md"
    export_markdown(sample_annotations, "test.pdf", out)
    assert "Great point!" in out.read_text()


def test_markdown_contains_pdf_name(sample_annotations, tmp_path):
    out = tmp_path / "out.md"
    export_markdown(sample_annotations, "test.pdf", out)
    assert "test.pdf" in out.read_text()


def test_markdown_empty_input(tmp_path):
    out = tmp_path / "empty.md"
    export_markdown([], "empty.pdf", out)
    assert out.exists()
    assert "# PDF Annotations" in out.read_text()


# ── export_plain_text ──────────────────────────────────────────────────────────

def test_plain_text_creates_file(sample_annotations, tmp_path):
    out = tmp_path / "out.txt"
    export_plain_text(sample_annotations, "test.pdf", out)
    assert out.exists() and out.stat().st_size > 0


def test_plain_text_contains_type(sample_annotations, tmp_path):
    out = tmp_path / "out.txt"
    export_plain_text(sample_annotations, "test.pdf", out)
    assert "Type    : Highlight" in out.read_text()


def test_plain_text_contains_marked_text(sample_annotations, tmp_path):
    out = tmp_path / "out.txt"
    export_plain_text(sample_annotations, "test.pdf", out)
    assert "Marked  : brown fox" in out.read_text()


def test_plain_text_contains_page_header(sample_annotations, tmp_path):
    out = tmp_path / "out.txt"
    export_plain_text(sample_annotations, "test.pdf", out)
    assert "PAGE 1" in out.read_text()


def test_plain_text_empty_input(tmp_path):
    out = tmp_path / "empty.txt"
    export_plain_text([], "empty.pdf", out)
    assert out.exists()
    assert "Total annotations: 0" in out.read_text()


# ── export_word ────────────────────────────────────────────────────────────────

def test_word_creates_file(sample_annotations, tmp_path):
    out = tmp_path / "out.docx"
    export_word(sample_annotations, "test.pdf", out)
    assert out.exists() and out.stat().st_size > 0


def test_word_contains_type_in_document(sample_annotations, tmp_path):
    from docx import Document
    out = tmp_path / "out.docx"
    export_word(sample_annotations, "test.pdf", out)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("Highlight" in t for t in texts)


def test_word_contains_pdf_name(sample_annotations, tmp_path):
    from docx import Document
    out = tmp_path / "out.docx"
    export_word(sample_annotations, "test.pdf", out)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("test.pdf" in t for t in texts)


# ── export_pdf ─────────────────────────────────────────────────────────────────

def test_pdf_creates_file(sample_annotations, tmp_path):
    out = tmp_path / "out.pdf"
    export_pdf(sample_annotations, "test.pdf", out)
    assert out.exists() and out.stat().st_size > 0


def test_pdf_has_magic_bytes(sample_annotations, tmp_path):
    out = tmp_path / "out.pdf"
    export_pdf(sample_annotations, "test.pdf", out)
    assert out.read_bytes()[:4] == b"%PDF"
